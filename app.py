import streamlit as st
import pandas as pd
from datetime import datetime
import math
import base64
import io
import time
import requests as req
from geopy.geocoders import Nominatim
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================================
# 1. GLOBAL UI & ELEGANT ENTERPRISE CSS THEME
# ==============================================================================
st.set_page_config(page_title="OptiStaff - Enterprise", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    :root{
        --bg:#F1F5F9; --surface:#FFFFFF; --border:#E6E9F0; --border-strong:#D6DBE6;
        --text:#0F172A; --muted:#64748B;
        --primary:#4F46E5; --primary-hover:#4338CA; --primary-soft:#EEF0FF;
        --accent:#06B6D4;
        --success:#10B981; --warning:#F59E0B; --danger:#EF4444;
        --radius:14px; --shadow:0 1px 2px rgba(15,23,42,.04), 0 10px 26px rgba(15,23,42,.06);
    }

    html, body, [class*="css"]{
        font-family:'Inter',-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;
        -webkit-font-smoothing:antialiased;
    }

    .stApp{
        background:
            radial-gradient(1100px 520px at 100% -8%, #EEF2FF 0%, rgba(238,242,255,0) 55%),
            radial-gradient(900px 460px at -10% 0%, #ECFEFF 0%, rgba(236,254,255,0) 52%),
            var(--bg);
    }

    @keyframes fadeIn { 0% { opacity: 0; transform: translateY(10px); } 100% { opacity: 1; transform: translateY(0); } }
    .main .block-container {
        max-width: 1140px;
        padding-top: 2.2rem;
        padding-bottom: 4rem;
        animation: fadeIn 0.5s ease-out forwards;
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header[data-testid="stHeader"] { background: transparent !important; box-shadow: none !important; }
    
    [data-testid="stSidebarCollapsedControl"], [data-testid="collapsedControl"] {
        visibility: visible !important; opacity: 1 !important; z-index: 999999 !important;
    }
    [data-testid="stSidebarCollapsedControl"] svg, [data-testid="collapsedControl"] svg,
    [data-testid="stSidebarCollapsedControl"] button, [data-testid="collapsedControl"] button {
        color: #4F46E5 !important; fill: #4F46E5 !important;
    }

    section[data-testid="stMain"], .main { color: #1E293B; }
    section[data-testid="stMain"] h1, .main h1 { color: #0F172A !important; font-weight: 700; letter-spacing: -0.02em; }
    section[data-testid="stMain"] h2, section[data-testid="stMain"] h3, section[data-testid="stMain"] h4,
    .main h2, .main h3, .main h4 { color: #0F172A !important; font-weight: 600; letter-spacing: -0.01em; }
    section[data-testid="stMain"] [data-testid="stMarkdownContainer"] p,
    .main [data-testid="stMarkdownContainer"] p { color: #1E293B; }
    
    [data-testid="stSidebar"] [data-testid="stCaptionContainer"] { color: #94A3B8 !important; }
    section[data-testid="stMain"] [data-testid="stCaptionContainer"] { color: var(--muted) !important; }

    div.stButton > button, div.stDownloadButton > button, div.stFormSubmitButton > button {
        border-radius: 10px;
        border: 1px solid transparent;
        font-weight: 600;
        background-color: var(--primary) !important;
        color: white !important;
        transition: all 0.18s cubic-bezier(0.25, 0.1, 0.25, 1);
        padding: 0.6rem 1.1rem;
        box-shadow: 0 1px 2px rgba(79,70,229,0.25);
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover, div.stFormSubmitButton > button:hover {
        background-color: var(--primary-hover) !important;
        transform: translateY(-1px);
        box-shadow: 0 8px 18px rgba(79, 70, 229, 0.28) !important;
    }
    div.stButton > button:active { transform: translateY(0) scale(0.99); }

    button[kind="secondary"] {
        background-color: var(--surface) !important;
        color: var(--text) !important;
        border: 1px solid var(--border-strong) !important;
        box-shadow: none !important;
    }
    button[kind="secondary"]:hover {
        background-color: #F8FAFC !important;
        border-color: var(--primary) !important;
        color: var(--primary) !important;
    }

    .stTextInput input, .stTextArea textarea, .stNumberInput input,
    .stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"] {
        border-radius: 10px !important;
        border-color: var(--border-strong) !important;
        background-color: var(--surface) !important;
        color: var(--text) !important;
        caret-color: var(--primary) !important;
        transition: border-color 0.2s, box-shadow 0.2s;
    }
    .stTextInput input, .stTextArea textarea, .stNumberInput input {
        -webkit-text-fill-color: var(--text) !important;
    }
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {
        color: #94A3B8 !important; -webkit-text-fill-color: #94A3B8 !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(79,70,229,0.15) !important;
    }
    label { font-weight: 500 !important; color: #334155 !important; }

    [data-testid="stForm"] {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.5rem 1.5rem 0.6rem;
        box-shadow: var(--shadow);
    }
    [data-testid="stExpander"] {
        border: 1px solid var(--border) !important;
        border-radius: 12px !important;
        background: var(--surface);
        box-shadow: var(--shadow);
        overflow: hidden;
    }

    .stTabs [data-baseweb="tab-list"] { gap: 0.35rem; border-bottom: 1px solid var(--border); overflow-x: auto; }
    .stTabs [data-baseweb="tab"] {
        height: auto; padding: 0.55rem 1rem; border-radius: 10px 10px 0 0;
        color: var(--muted); font-weight: 600;
    }
    .stTabs [aria-selected="true"] { color: var(--primary) !important; background: var(--primary-soft); }
    .stTabs [data-baseweb="tab-highlight"] { background-color: var(--primary); }

    div[role="radiogroup"] { gap: 0.4rem; }
    div[role="radiogroup"] label {
        background: var(--surface); border: 1px solid var(--border-strong); border-radius: 10px;
        padding: 0.35rem 0.85rem; transition: all 0.15s ease;
    }
    div[role="radiogroup"] label:hover { border-color: var(--primary); }

    [data-testid="stMetric"] {
        background: var(--surface); border: 1px solid var(--border); border-radius: 12px;
        padding: 1rem 1.1rem; box-shadow: var(--shadow);
    }
    [data-testid="stMetricValue"] { color: var(--text); font-weight: 700; }

    [data-testid="stAlert"] { border-radius: 12px; }
    [data-testid="stDataFrame"] { border: 1px solid var(--border); border-radius: 12px; overflow: hidden; }

    hr { margin: 1.1rem 0; border-color: var(--border); }
    img { max-width: 100%; height: auto; border-radius: 12px; }

    [data-testid="stSidebar"] { background: #0F172A; border-right: 1px solid #1E293B; }
    [data-testid="stSidebar"] * { color: #E2E8F0; }
    [data-testid="stSidebar"] .stButton > button {
        background-color: #1E293B !important; border: 1px solid #334155 !important;
        color: #E2E8F0 !important; box-shadow: none !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background-color: #334155 !important; color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] [data-testid="stAlert"] {
        background: rgba(255,255,255,0.06); border: 1px solid rgba(255,255,255,0.10);
    }

    .os-sb-brand { display: flex; align-items: center; gap: 0.55rem; font-size: 1.25rem; font-weight: 700; color: #FFFFFF !important; padding: 0.2rem 0 0.1rem; }
    .os-sb-brand .dot { width: 26px; height: 26px; border-radius: 7px; background: linear-gradient(135deg, #6366F1, #22D3EE); }

    .os-hero { text-align: center; margin: 0.5rem auto 1.6rem; }
    .os-logo { display: inline-flex; align-items: center; gap: 0.6rem; font-weight: 700; font-size: 1.7rem; color: var(--text); }
    .os-logo .dot { width: 34px; height: 34px; border-radius: 9px; background: linear-gradient(135deg, #4F46E5, #06B6D4); display: inline-block; box-shadow: 0 6px 16px rgba(79,70,229,.35); }
    .os-tag { color: var(--muted); margin-top: 0.5rem; font-size: 0.95rem; }

    @media (max-width: 640px) {
        .main .block-container { max-width: 100% !important; padding-left: 1rem; padding-right: 1rem; padding-top: 1.2rem; }
        .os-hero { margin-top: 0.2rem; }
        .os-logo { font-size: 1.4rem; }
        .os-logo .dot { width: 28px; height: 28px; }
        h1 { font-size: 1.6rem !important; }
        h2, h3 { font-size: 1.15rem !important; }
        .stTabs [data-baseweb="tab"] { padding: 0.45rem 0.6rem; font-size: 0.9rem; }
        [data-testid="stForm"] { padding: 1.1rem 1.1rem 0.4rem; }
        [data-testid="stMetricValue"] { font-size: 1.4rem !important; }
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. STATE MEMORY INITIALIZATION
# ==============================================================================
for k, v in [('logged_in', False), ('user_role', None), ('user_name', None), ('user_event_id', None)]:
    if k not in st.session_state: 
        st.session_state[k] = v

# ==============================================================================
# 3. SUPABASE CONFIGURATION & REST API WRAPPER
# ==============================================================================
SUPA_URL = st.secrets["SUPABASE_URL"]
SUPA_KEY = st.secrets["SUPABASE_KEY"]

def _h(prefer=None):
    h = {
        "apikey": SUPA_KEY,
        "Authorization": f"Bearer {SUPA_KEY}",
        "Content-Type": "application/json"
    }
    if prefer: 
        h["Prefer"] = prefer
    return h

def sread(table, filters=None, order=None, single=False):
    try:
        p = dict(filters) if filters else {}
        if order: 
            p["order"] = order
        r = req.get(f"{SUPA_URL}/rest/v1/{table}", headers=_h(), params=p, timeout=15)
        r.raise_for_status()
        d = r.json()
        return d[0] if (single and d) else d
    except Exception:
        return None if single else []

def sins(table, data):
    try:
        r = req.post(f"{SUPA_URL}/rest/v1/{table}", headers=_h("return=minimal"), json=data, timeout=15)
        return r.status_code in [200, 201]
    except Exception as e: 
        st.error(f"Kesalahan Basis Data (Insert): {e}")
        return False

def sups(table, data):
    try:
        r = req.post(f"{SUPA_URL}/rest/v1/{table}", headers=_h("resolution=merge-duplicates,return=minimal"), json=data, timeout=15)
        return r.status_code in [200, 201]
    except Exception as e: 
        st.error(f"Kesalahan Basis Data (Upsert): {e}")
        return False

def supd(table, filters, data):
    try:
        p = {k: f"eq.{v}" for k, v in filters.items()}
        r = req.patch(f"{SUPA_URL}/rest/v1/{table}", headers=_h("return=minimal"), params=p, json=data, timeout=15)
        return r.status_code in [200, 204]
    except Exception as e: 
        st.error(f"Kesalahan Basis Data (Update): {e}")
        return False

def todf(rows):
    if not rows: 
        return pd.DataFrame()
    df = pd.DataFrame(rows)
    for c in df.select_dtypes(include='object').columns: 
        df[c] = df[c].fillna('').astype(str)
    return df

# ==============================================================================
# 4. DOMAIN LOGIC FUNCTIONS
# ==============================================================================
def ev_all(status=None):
    return todf(sread("events", {"status": f"eq.{status}"} if status else None, "event_id.asc"))

def ev_save(eid, nm, lok, rad, div):
    sups("events", {
        "event_id": str(eid),
        "nama_event": str(nm),
        "lokasi_target": str(lok),
        "radius_meter": float(rad),
        "status": "Aktif",
        "status_lpj": "Belum Diserahkan",
        "divisions": str(div)
    })

def ev_divs(eid):
    r = sread("events", {"event_id": f"eq.{eid}"}, single=True)
    if r and r.get('divisions'): 
        return [d.strip() for d in str(r['divisions']).split(',') if d.strip()]
    return ["Ticketing", "Logistik"]

def ev_lpj(eid): 
    supd("events", {"event_id": eid}, {"status_lpj": "Sudah Diserahkan"})

def ev_close(eid): 
    supd("events", {"event_id": eid}, {"status": "Selesai"})
    supd("crews", {"event_id": eid}, {"status": "Expired"})

def cr_all(eid=None, role=None, div=None):
    f = {}
    if eid: f["event_id"] = f"eq.{eid}"
    if role: f["role"] = f"eq.{role}"
    if div: f["division"] = f"eq.{div}"
    return todf(sread("crews", f if f else None, "nama.asc"))

def cr_save(cid, nm, em, role, eid, pw, div):
    sups("crews", {
        "crew_id": str(cid),
        "nama": str(nm),
        "email": str(em),
        "role": str(role),
        "event_id": str(eid),
        "password": str(pw),
        "status": "Aktif",
        "division": str(div)
    })

def att_all(eid=None):
    return todf(sread("attendances", {"event_id": f"eq.{eid}"} if eid else None, "timestamp.desc"))

def att_save(aid, eid, cn, ts, lat, lon, sg, sk, foto, tipe):
    sins("attendances", {
        "attendance_id": str(aid),
        "event_id": str(eid),
        "crew_name": str(cn),
        "timestamp": str(ts),
        "latitude": float(lat),
        "longitude": float(lon),
        "status_geotag": str(sg),
        "status_pekerjaan": str(sk),
        "foto_selfie_url": str(foto),
        "tipe_absen": str(tipe)
    })

def tk_all(event_id=None, eid=None, div=None, cn=None, status=None):
    if event_id is not None and eid is None:
        eid = event_id
    filters = {}
    if eid:
        filters["event_id"] = f"eq.{eid}"
    if div:
        filters["division"] = f"eq.{div}"
    if cn:
        filters["crew_name"] = f"eq.{cn}"
    if status:
        filters["status"] = f"eq.{status}"
    rows = sread(
        "tasks",
        filters if filters else None,
        "created_at.desc"
    )
    return todf(rows)

def tk_create(eid, div, cn, tn, ref=""):
    tid = f"TSK-{int(time.time()*1000)}"
    sins("tasks", {
        "task_id": tid,
        "event_id": str(eid),
        "division": str(div),
        "crew_name": str(cn),
        "task_name": str(tn),
        "status": "Belum Selesai",
        "photo_report": "",
        "reference_photo": str(ref),
        "tl_comment": "",
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })
    return tid

def tk_upd(tid, data): 
    supd("tasks", {"task_id": tid}, data)

def pg_latest(tid):
    r = sread("task_progress", {"task_id": f"eq.{tid}"}, "timestamp.desc", single=False)
    return r[0] if r else None

def pg_save(tid, eid, cn, photo):
    pid = f"PROG-{tid}-{int(time.time())}"
    sins("task_progress", {
        "prog_id": pid,
        "task_id": str(tid),
        "event_id": str(eid),
        "crew_name": str(cn),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "photo_report": str(photo),
        "status": "Menunggu Verifikasi",
        "tl_comment": ""
    })
    tk_upd(tid, {"status": "Menunggu Verifikasi"})

def tk_eval(tid, status, kom=""):
    tk_upd(tid, {"status": status, "tl_comment": kom})
    p = pg_latest(tid)
    if p: 
        supd("task_progress", {"prog_id": p['prog_id']}, {"status": status, "tl_comment": kom})

def foto_compress(f):
    raw = Image.open(io.BytesIO(f.getvalue()))
    if raw.mode in ("RGBA", "P"): 
        raw = raw.convert("RGB")
    raw.thumbnail((400, 400))
    buf = io.BytesIO()
    raw.save(buf, format="JPEG", quality=60)
    return "data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode()

def foto_show(val, cap="Lampiran Visual", use_col_width=True):
    if not val or len(str(val)) < 10: 
        return
    v = str(val).strip()
    if v.startswith("http"): 
        st.image(v, caption=cap, use_container_width=use_col_width)
    elif "base64," in v:
        try: 
            st.image(Image.open(io.BytesIO(base64.b64decode(v.split("base64,")[1]))), caption=cap, use_container_width=use_col_width)
        except Exception: 
            pass

def jarak(lat1, lon1, lat2, lon2):
    R = 6371000
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dp, dl = math.radians(lat2 - lat1), math.radians(lon2 - lon1)
    a = math.sin(dp / 2)**2 + math.cos(p1) * math.cos(p2) * math.sin(dl / 2)**2
    return R * (2 * math.atan2(math.sqrt(a), math.sqrt(1 - a)))

def fixk(val, is_lat=True):
    try:
        s = f"{val}".replace(',', '').replace('.', '').strip()
        neg = s.startswith('-')
        if neg: s = s[1:]
        n = float(s)
        if is_lat:
            while n > 9.0: n /= 10.0
            return -n if neg else n
        else:
            while n > 180.0: n /= 10.0
            if n < 90.0:
                while n < 100.0: n *= 10.0
            return -n if neg else n
    except: 
        return -6.2181 if is_lat else 106.8024

def get_location_name(coord):
    try:
        lat, lon = map(float, coord.split(","))
        geolocator = Nominatim(user_agent="optistaff")
        location = geolocator.reverse((lat, lon), language="id")
        if location:
            return location.address
    except:
        pass
    return coord

def mk_lpj(df_att, eid, nm_ev, lokasi_event, df_tk):
    doc = Document()
    for s in doc.sections: 
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("OPTI STAFF CLOUD SYSTEM\nExecutive Report Management")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(140, 140, 142)
    
    r2 = doc.add_paragraph().add_run("Laporan Pertanggungjawaban")
    r2.font.size = Pt(24)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(29, 29, 31)
    
    doc.add_paragraph().add_run(f"Project ID: {eid} — {nm_ev}")
    doc.add_paragraph().add_run(f"Nama Event : {nm_ev}")
    doc.add_paragraph().add_run(f"Lokasi Event : {lokasi_event}")
    doc.add_paragraph().add_run(f"Tanggal Generate : {datetime.now().strftime('%d %B %Y %H:%M')}")

    doc.add_paragraph().add_run("—" * 65)
    doc.add_paragraph().add_run("—" * 65)
    
    doc.add_heading("1. Ringkasan Operasional Lapangan", level=1)
    p2 = doc.add_paragraph()
    tot = len(df_att)
    val = len(df_att[df_att['status_geotag'].astype(str).str.contains("MATCH", na=False)]) if not df_att.empty and 'status_geotag' in df_att.columns else 0
    p2.add_run(f"Waktu Penarikan Dokumen: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nTotal Aktivitas: {tot} Transaksi | Valid: {val} | Indikasi Pelanggaran: {tot - val}\n")
    
    if not df_tk.empty:
        doc.add_heading("2. Status Capaian Tugas Operasional", level=1)
        tt = len(df_tk)
        td = len(df_tk[df_tk['status'] == 'Disetujui'])
        doc.add_paragraph().add_run(f"Total Tugas: {tt} | Diselesaikan: {td} | Menunggu Tindakan: {tt - td}")
    
    doc.add_heading("3. Log Detail Presensi", level=1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Nama", "Waktu", "Tipe", "Geotag", "Progress"]):
        tbl.rows[0].cells[i].text = h
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
    for _, row in df_att.iterrows():
        rc = tbl.add_row().cells
        rc[0].text = str(row.get('crew_name', ''))
        rc[1].text = str(row.get('timestamp', ''))
        rc[2].text = str(row.get('tipe_absen', ''))
        rc[3].text = str(row.get('status_geotag', ''))
        rc[4].text = str(row.get('status_pekerjaan', ''))
        
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ==============================================================================
# 5. GATEWAY LOGIN
# ==============================================================================
if not st.session_state.logged_in:
    st.markdown("""
    <div class="os-hero">
        <div class="os-logo"><span class="dot"></span> OptiStaff</div>
        <div class="os-tag">Sistem Manajemen Kehadiran Terintegrasi &middot; Geofencing &amp; Database Divisi</div>
    </div>
    """, unsafe_allow_html=True)
    
    _, mid, _ = st.columns([1, 1.6, 1])
    with mid:
        with st.form("lg"):
            em = st.text_input("Email Pengguna")
            pw = st.text_input("Kata Sandi", type="password")
            
            if st.form_submit_button("Masuk ke Sistem", use_container_width=True):
                with st.spinner("Memvalidasi kredensial..."):
                    if em == "eo@optistaff.com" and pw == "adminpro":
                        st.session_state.update({
                            "logged_in": True, 
                            "user_role": "EO (Organizer)", 
                            "user_name": "Administrator Pusat", 
                            "user_event_id": "ALL"
                        })
                        time.sleep(.3)
                        st.rerun()
                    else:
                        rows = sread("crews", {"email": f"eq.{em}", "status": "eq.Aktif"})
                        u = rows[0] if rows else None
                        if u:
                            if str(u['password']) == str(pw):
                                st.session_state.update({
                                    "logged_in": True, 
                                    "user_role": u['role'], 
                                    "user_name": u['nama'], 
                                    "user_event_id": u['event_id']
                                })
                                time.sleep(.3)
                                st.rerun()
                            else: 
                                st.error("Kata sandi tidak valid.")
                        else: 
                            st.error("Email tidak ditemukan atau status akun tidak aktif.")
    st.stop()

# ==============================================================================
# 6. SIDEBAR NAVIGASI UMUM
# ==============================================================================
st.sidebar.markdown('<div class="os-sb-brand"><span class="dot"></span> OptiStaff</div>', unsafe_allow_html=True)
st.sidebar.caption("Operations Control System")
st.sidebar.markdown("---")
st.sidebar.markdown(f"**Pengguna:** {st.session_state.user_name}")
st.sidebar.info(f"Hak Akses: {st.session_state.user_role}")

if st.session_state.user_event_id and st.session_state.user_event_id != "ALL": 
    st.sidebar.success(f"ID Proyek: {st.session_state.user_event_id}")
else: 
    st.sidebar.warning("Mode: Akses Administrator")

st.sidebar.markdown("---")
if st.sidebar.button("Segarkan Data Server", use_container_width=True): 
    st.rerun()
if st.sidebar.button("Keluar Sistem", type="secondary", use_container_width=True): 
    st.session_state.logged_in = False
    st.rerun()

# ==============================================================================
# 7. INTERFACE ROLE: EO (EVENT ORGANIZER)
# ==============================================================================
if st.session_state.user_role == "EO (Organizer)":
    st.title("Portal Manajemen Pusat")
    t1, t2, t3 = st.tabs(["Manajemen Proyek", "Registrasi Personel", "Penutupan Proyek"])
    
    with t1:
        st.subheader("Pendaftaran Proyek & Pembuatan Divisi Operasional")
        df_ev = ev_all()
        c1, c2 = st.columns(2)
        c1.metric("Total Proyek Terdaftar", len(df_ev))
        c2.metric("Proyek Berjalan Aktif", len(df_ev[df_ev['status'] == 'Aktif']) if not df_ev.empty else 0)
        
        with st.form("fev", clear_on_submit=True):
            eid = st.text_input("ID Proyek Baru:", value=f"EVT0{len(df_ev) + 1}")
            enm = st.text_input("Nama Proyek:")
            elok = st.text_input("Koordinat Pusat (Lat, Lon):", value="-6.2181, 106.8024")
            erad = st.slider("Radius Geofence Keamanan (Meter):", 50, 500, 100)
            
            st.markdown("**Konfigurasi Divisi Eksekusi Lapangan**")
            ddef = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
            dsel = st.multiselect("Struktur Divisi Standar Sistem:", ddef, default=ddef)
            dkust = st.text_input("Divisi Kustom Tambahan (Pisahkan dengan koma):", value="")
            
            if st.form_submit_button("Aktifkan Proyek & Divisi", use_container_width=True) and enm:
                adiv = dsel + [d.strip() for d in dkust.split(",") if d.strip()]
                ev_save(eid, enm, elok, erad, ",".join(adiv))
                st.success(f"Proyek '{enm}' berhasil diinisiasi.")
                time.sleep(1)
                st.rerun()
                
        st.markdown("**Daftar Status Proyek Aktif**")
        st.dataframe(df_ev, use_container_width=True)
        
    with t2:
        st.subheader("Distribusi Personel & Pemetaan Divisi")
        df_cr = cr_all()
        df_evl = ev_all()
        evlist = df_evl["event_id"].tolist() if not df_evl.empty else ["EVT01"]
        
        st.markdown("**1. Pilih Proyek Penugasan Terlebih Dahulu:**")
        selv = st.selectbox("Target Proyek:", evlist, label_visibility="collapsed")
        dopts = ev_divs(selv)
        
        st.markdown("**2. Konfigurasi Hak Akses Personel:**")
        crole = st.radio("Tingkat Akses:", ["Team Leader", "Crew"], horizontal=True)
        
        with st.form("fcr", clear_on_submit=True):
            cid = st.text_input("ID Personel:", value=f"CRW0{len(df_cr) + 1}")
            cnm = st.text_input("Nama Lengkap Sesuai KTP:")
            cem = st.text_input("Email Login Akun:")
            cpw = st.text_input("Kata Sandi Akun:", value="12345")
            
            if crole == "Crew": 
                cdiv = st.selectbox("Alokasi Divisi Operasional (Spesifik):", dopts)
            else: 
                st.info("Otoritas Team Leader bersifat absolut dan mencakup seluruh divisi pada proyek ini.")
                cdiv = "Team Leader"
                
            if st.form_submit_button("Daftarkan Personel Baru", use_container_width=True) and cnm and cem:
                cr_save(cid, cnm, cem, crole, selv, cpw, cdiv)
                st.success(f"Kredensial Personel {cnm} terdaftar sukses sebagai {crole} pada divisi: {cdiv}.")
                time.sleep(1)
                st.rerun()
                
        st.markdown("**Direktori Registrasi Personel**")
        st.dataframe(df_cr, use_container_width=True)
        
    with t3:
        st.subheader("Penyelesaian & Terminasi Dokumen Proyek")
        df_eva = ev_all(status='Aktif')
        if df_eva.empty: 
            st.info("Tidak ada proyek aktif saat ini.")
        else:
            c1, c2 = st.columns([3, 1])
            ecid = c1.selectbox("Pilih Proyek:", df_eva["event_id"].tolist(), key="sec")
            einfo = df_eva[df_eva['event_id'] == ecid].iloc[0]
            st.info(f"Status Otorisasi Dokumen LPJ: {einfo.get('status_lpj', 'Belum Diserahkan')}")
            
            if c2.button("Tutup Paksa Proyek", type="primary", use_container_width=True):
                if einfo.get('status_lpj') != "Sudah Diserahkan": 
                    st.error("Proses Terminasi Ditolak. LPJ belum diserahkan oleh Pengawas (Team Leader).")
                else: 
                    ev_close(ecid)
                    st.success(f"Proyek {ecid} resmi ditutup.")
                    time.sleep(1)
                    st.rerun()
                    
            st.markdown("---")
            st.subheader("Riwayat Evaluasi Log Kehadiran Proyek")
            df_attm = att_all(eid=ecid)
            if df_attm.empty: 
                st.info("Data kehadiran belum tersedia.")
            else:
                cols = [c for c in ['attendance_id', 'crew_name', 'tipe_absen', 'timestamp', 'status_geotag', 'status_pekerjaan'] if c in df_attm.columns]
                st.dataframe(df_attm[cols], use_container_width=True)

# ==============================================================================
# 8. INTERFACE ROLE: TEAM LEADER (PENGAWAS OPERASIONAL)
# ==============================================================================
elif st.session_state.user_role == "Team Leader":
    eid = st.session_state.user_event_id
    st.title("Dashboard Pengawas Operasional")
    st.caption(f"ID Proyek Penugasan Aktif: {eid}")
    
    tm, td, ta = st.tabs(["Pemantauan Radar", "Manajemen Tugas Divisi", "Manajemen Personel Lokal"])
    divs = ev_divs(eid)
    
    with tm:
        df_att = att_all(eid=eid)
        evr = sread("events", {"event_id": f"eq.{eid}"}, single=True) or {}
        
        st.subheader("Log Kehadiran Personel Aktual")
        if df_att.empty: 
            st.info("Sistem belum mendeteksi adanya data kehadiran pada proyek ini.")
        else:
            tot = len(df_att)
            val = len(df_att[df_att['status_geotag'].astype(str).str.contains("MATCH", na=False)]) if 'status_geotag' in df_att.columns else 0
            
            c1, c2, c3 = st.columns(3)
            c1.metric("Total Percobaan Kehadiran", tot)
            c2.metric("Validasi Geofence Sesuai", val)
            c3.metric("Pelanggaran Batas Radar", tot - val)
            
            cols = [c for c in df_att.columns if c != 'foto_selfie_url']
            st.dataframe(df_att[cols], use_container_width=True)
            
            st.markdown("---")
            st.subheader("Radar Spasial Geolocation")
            if {'latitude', 'longitude'}.issubset(df_att.columns):
                dm = df_att[['latitude', 'longitude']].copy()
                dm['latitude'] = dm['latitude'].apply(lambda x: fixk(x, True))
                dm['longitude'] = dm['longitude'].apply(lambda x: fixk(x, False))
                dm = dm.dropna()
                if not dm.empty: 
                    st.map(dm, latitude='latitude', longitude='longitude', size=40)
                
            st.markdown("---")
            st.subheader("Otorisasi Dokumen Laporan Akhir (LPJ)")

            dftk = tk_all(eid=eid)

            koordinat = evr.get("lokasi_target", "")
            lokasi_event = f"{koordinat}\n({get_location_name(koordinat)})"

            docx = mk_lpj(df_att, eid, evr.get("nama_event", ""), lokasi_event, dftk)

            plpj = dftk[dftk["status"] != "Disetujui"] if not dftk.empty else pd.DataFrame()
            
            cd1, cd2 = st.columns(2)
            with cd1: 
                st.download_button("Unduh Evaluasi Draf Dokumen (.docx)", data=docx, file_name=f"LPJ_{eid}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="secondary", use_container_width=True)
            with cd2:
                if not plpj.empty: 
                    st.error("Otorisasi LPJ Terkunci: Harap selesaikan dan setujui seluruh tugas divisi.")
                elif st.button("Sahkan Dokumen & Transmisikan ke Pusat", type="primary", use_container_width=True): 
                    ev_lpj(eid)
                    st.success("Tanda tangan digital berhasil. Dokumen LPJ telah diotorisasi dan diserahkan ke sistem pusat.")
                    
            st.markdown("---")
            st.subheader("Verifikasi Visual Kehadiran Individu")
            if not df_att.empty and 'crew_name' in df_att.columns:
                crew_list = sorted(df_att["crew_name"].dropna().unique().tolist())
                selected_crew = st.selectbox("Pilih Personel untuk Verifikasi:", crew_list, key="attendance_history")

                history = df_att[df_att["crew_name"] == selected_crew].sort_values("timestamp", ascending=False)

                for _, row in history.iterrows():
                    with st.expander(f"{row['tipe_absen']} | {row['timestamp']}", expanded=False):
                        c_img, c_txt = st.columns([1, 3])
                        with c_img:
                            foto_show(str(row.get("foto_selfie_url","")), "Citra Kehadiran")
                        with c_txt:
                            sg = str(row.get("status_geotag",""))
                            if "OUT OF RANGE" in sg:
                                st.error(f"Pelanggaran Lokasi Spasial: {sg}")
                            elif "MATCH" in sg:
                                st.success(f"Lokasi Valid Sesuai Koordinat: {sg}")
                            else:
                                st.info(f"Status Geofence: {sg}")

                            st.write(f"**Jenis Kehadiran:** {row.get('tipe_absen','-')}")
                            st.write(f"**Laporan Pekerjaan:** {row.get('status_pekerjaan','-')}")
            else:
                st.info("Belum ada data kehadiran untuk diverifikasi.")
                    
    with td:
        st.subheader("Distribusi Tugas Operasional")
        tdiv = st.selectbox("Pilih Fokus Divisi Operasional:", divs)
        dfk = cr_all(eid=eid, role='Crew', div=tdiv)
        krul = dfk['nama'].tolist() if not dfk.empty else []
        
        st.markdown("---")
        st.subheader("1. Alokasi Tugas Baru")
        if not krul: 
            st.info("Sistem belum mendeteksi adanya personel aktif pada divisi ini.")
        else:
            st.caption("Daftar Personel Tersedia: " + ", ".join(krul))
            with st.form("ftk", clear_on_submit=True):
                ttxt = st.text_area("Parameter Instruksi Tugas (1 baris teks mewakili 1 instruksi kerja utuh):", value="Melakukan verifikasi alat scanner barcode di Gate A")
                
                with st.expander("Gunakan Kamera / Tambah Foto Referensi"):
                    rcam = st.camera_input("Ambil Foto Referensi")
                
                if st.form_submit_button("Alokasikan Tugas kepada Personel", use_container_width=True):
                    lines = [l.strip() for l in ttxt.split("\n") if l.strip()]
                    if lines:
                        rurl = foto_compress(rcam) if rcam else ""
                        with st.spinner("Memproses transmisi data instruksi..."):
                            for i, tn in enumerate(lines): 
                                tk_create(eid, tdiv, krul[i % len(krul)], tn, rurl)
                        st.success("Seluruh tugas berhasil dialokasikan secara berurutan kepada personel terkait.")
                        time.sleep(1)
                        st.rerun()
                        
        st.markdown("---")
        st.subheader("2. Verifikasi Laporan Tugas Personel")
        pend = tk_all(eid=eid, div=tdiv, status="Menunggu Verifikasi") 
        
        if pend.empty: 
            st.info("Pusat verifikasi bersih. Tidak ada laporan baru yang menunggu tinjauan Anda.")
        else:
            for _, row in pend.iterrows():
                with st.expander(f"Tinjauan Laporan Personel: {row['crew_name']} | Tugas: {row['task_name']}", expanded=True):
                    pg = pg_latest(row['task_id'])
                    ph = pg['photo_report'] if pg else ""
                    if ph: 
                        foto_show(ph, "Lampiran Visual Penunjang")
                    else: 
                        st.info("Personel tidak menyertakan foto, atau arsip visual tidak ditemukan dalam server.")
                        
                    kom = st.text_input("Komentar Evaluasi Pengawas (Wajib diisi untuk tindakan Penolakan/Revisi):", key=f"k_{row['task_id']}")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Setujui Laporan", key=f"app_{row['task_id']}", use_container_width=True):
                            tk_eval(row['task_id'], "Disetujui", kom)
                            time.sleep(1)
                            st.rerun()
                    with c2:
                        if st.button("Tolak Laporan (Kirim Ulang)", key=f"rej_{row['task_id']}", use_container_width=True):
                            if not kom: 
                                st.error("Sistem Audit Menolak: Harap sertakan komentar sebagai panduan perbaikan.")
                            else: 
                                tk_eval(row['task_id'], "Ditolak", kom)
                                time.sleep(1)
                                st.rerun()
                                
        st.markdown("---")
        st.subheader("3. Papan Status Kinerja Keseluruhan Divisi")
        atk = tk_all(eid=eid, div=tdiv) 
        if atk.empty: 
            st.info("Data riwayat pekerjaan operasional belum tersedia di divisi ini.")
        else: 
            st.dataframe(atk[['crew_name', 'task_name', 'status', 'tl_comment']], use_container_width=True)

        st.markdown("---")
        st.subheader("4. Riwayat Laporan Tugas per Personel")
        if atk.empty:
            st.info("Data arsip pekerjaan belum tersedia untuk dieksplorasi.")
        else:
            c_names = atk['crew_name'].dropna().unique().tolist()
            sel_cname = st.selectbox("Pilih Nama Personel untuk Meninjau Arsip Tugas:", sorted(c_names))
            c_tasks = atk[atk['crew_name'] == sel_cname]
            st.dataframe(c_tasks[['task_name', 'status', 'tl_comment', 'created_at']], use_container_width=True)
            
    with ta:
        st.subheader("Rekrutmen Mandiri Personel Lokal (Akses Pengawas)")
        with st.form("fta", clear_on_submit=True):
            nid = st.text_input("ID Personel Darurat:", value=f"CRW-{datetime.now().strftime('%H%M')}")
            nnm = st.text_input("Nama Lengkap:")
            nem = st.text_input("Email Login Akun:")
            npw = st.text_input("Kata Sandi (Password):", value="12345")
            ndiv = st.selectbox("Alokasi Peminjaman Divisi Operasional:", divs)
            
            if st.form_submit_button("Otorisasi & Daftarkan Personel Baru", type="primary", use_container_width=True) and nnm and nem:
                cr_save(nid, nnm, nem, "Crew", eid, npw, ndiv)
                st.success("Data kredensial personel berhasil diregistrasikan ke dalam sistem otorisasi pusat.")
                time.sleep(1)
                st.rerun()

# ==============================================================================
# 9. INTERFACE ROLE: CREW (PERSONEL EKSEKUTOR LAPANGAN)
# ==============================================================================
else:
    eid = st.session_state.user_event_id
    st.title("Portal Operasional Lapangan")
    
    crrows = sread("crews", {"nama": f"eq.{st.session_state.user_name}", "event_id": f"eq.{eid}"})
    crinfo = crrows[0] if crrows else {}
    mydiv = crinfo.get('division', 'Umum')
    
    st.caption(f"Personel Bertugas: {st.session_state.user_name} | Alokasi Divisi: {mydiv}")
    st.markdown("---")
    
    # CEK STATUS ABSEN MASUK HARI INI
    today_str = datetime.now().strftime("%Y-%m-%d")
    df_absen = att_all(eid=eid)
    
    my_absen_today = pd.DataFrame()
    if not df_absen.empty and 'crew_name' in df_absen.columns and 'timestamp' in df_absen.columns:
        my_absen_today = df_absen[
            (df_absen["crew_name"] == st.session_state.user_name) & 
            (df_absen["timestamp"].str.startswith(today_str))
        ]

    sudah_masuk = False
    if not my_absen_today.empty:
        valid_masuk = my_absen_today[
            (my_absen_today["tipe_absen"].astype(str).str.contains("Masuk", case=False)) & 
            (my_absen_today["status_geotag"].astype(str).str.contains("MATCH", case=False))
        ]
        if not valid_masuk.empty:
            sudah_masuk = True

    # LOGIKA MENU DINAMIS BERDASARKAN STATUS ABSEN MASUK
    if sudah_masuk:
        st.info("Akses operasional dibuka. Anda telah menyelesaikan proses Absen Masuk secara valid.")
        menu_options = ["Absen Pulang", "To-Do List Divisi"]
    else:
        st.warning("Silakan selesaikan Absen Masuk di dalam radius area untuk membuka akses operasional kerja.")
        menu_options = ["Absen Masuk"]

    menu = st.radio("Pilih Modul Operasional:", menu_options, horizontal=True)
    
    if menu in ["Absen Masuk", "Absen Pulang"]:
        st.subheader(f"Validasi Presensi Aktif ({menu})")
        
        allow_absen = True
        if menu == "Absen Pulang":
            my_tasks_eval = tk_all(eid=eid, cn=st.session_state.user_name)
            if not my_tasks_eval.empty and 'status' in my_tasks_eval.columns:
                pending_eval = my_tasks_eval[my_tasks_eval['status'] != 'Disetujui']
                if not pending_eval.empty:
                    allow_absen = False
                    st.error("Otorisasi Absen Pulang Terkunci: Terdapat instruksi tugas yang belum diselesaikan atau sedang menunggu verifikasi Pengawas.")
                    st.dataframe(pending_eval[['task_name', 'status']], use_container_width=True, hide_index=True)

        if allow_absen:
            evr = sread("events", {"event_id": f"eq.{eid}"})
            evd = evr[0] if evr else {}
            lok = evd.get('lokasi_target', "-6.2181, 106.8024")
            rad = float(evd.get('radius_meter', 100))
            
            foto = st.camera_input("Pengambilan Otorisasi Wajah Real-Time")
            c1, c2 = st.columns(2)
            lat = c1.number_input("Latitude Saat Ini:", value=-6.2181, format="%.5f")
            lon = c2.number_input("Longitude Saat Ini:", value=106.8024, format="%.5f")
            lap = st.text_area("Laporan Kondisi Lapangan (Catatan Operasional Shift):")
            
            if st.button(f"Kirim Data Otorisasi {menu}", type="primary", use_container_width=True) and foto:
                with st.spinner("Memproses transmisi data presensi ke server pusat..."):
                    curl = foto_compress(foto)
                    try: 
                        lt, ln = map(float, str(lok).split(","))
                    except: 
                        lt, ln = -6.2181, 106.8024
                    
                    latc = fixk(lat, True)
                    lonc = fixk(lon, False)
                    j = jarak(latc, lonc, lt, ln)
                    sg = "MATCH (Valid)" if j <= rad else "OUT OF RANGE (Fraud Suspect)"
                    
                    att_save(f"ATT-{datetime.now().strftime('%H%M%S')}", eid, st.session_state.user_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), latc, lonc, sg, lap, curl, menu)
                    
                if j <= rad: 
                    st.success(f"Otorisasi Absensi Tervalidasi. Jarak ke pusat operasional: {j:.1f} m.")
                else: 
                    st.warning(f"Otorisasi Ditolak. Posisi berada di luar batas keamanan yang diizinkan ({j:.1f} m).")
                time.sleep(2)
                st.rerun()
            
    elif menu == "To-Do List Divisi":
        st.subheader("Lembar Kerja Eksekusi & Verifikasi Laporan")
        allmy = tk_all(eid=eid, cn=st.session_state.user_name) 
        
        todo = allmy[allmy['status'] == "Belum Selesai"] if not allmy.empty else pd.DataFrame()
        rej = allmy[allmy['status'] == "Ditolak"] if not allmy.empty else pd.DataFrame()
        pend = allmy[allmy['status'] == "Menunggu Verifikasi"] if not allmy.empty else pd.DataFrame()
        appr = allmy[allmy['status'] == "Disetujui"] if not allmy.empty else pd.DataFrame()
        
        st.markdown("### Daftar Instruksi Tugas Prioritas")
        if todo.empty: 
            st.info("Papan intruksi bersih. Tidak ada tugas alokasi baru dari Pengawas Divisi.")
        else:
            for _, row in todo.iterrows():
                with st.container():
                    st.markdown(f"**Detail Tugas:** {row['task_name']}")
                    
                    ref_photo = str(row.get("reference_photo", "")).strip()

                    if len(ref_photo) > 10:

                        c_img, c_inf = st.columns([1,4])

                        with c_img:
                            foto_show(
                                ref_photo,
                                "📷 Referensi dari Team Leader"
                            )

                        with c_inf:
                            st.info(
                                "Foto ini merupakan referensi pekerjaan yang dikirim oleh Team Leader."
                            )

                    cam = st.camera_input("Lampirkan Bukti Penyelesaian", key=f"cam_{row['task_id']}")
                    
                    if st.button(f"Kirim Laporan Penyelesaian: {row['task_name']}", key=f"btn_{row['task_id']}", use_container_width=True):
                        if not cam: 
                            st.error("Protokol Audit: Harap lampirkan bukti visual penyelesaian kerja untuk pelaporan.")
                        else:
                            with st.spinner("Mengunggah dokumen bukti penyelesaian ke server..."): 
                                url = foto_compress(cam)
                                pg_save(row['task_id'], eid, st.session_state.user_name, url)
                            st.success("Laporan berhasil diserahkan ke sistem antrean verifikasi Pengawas.")
                            time.sleep(1.5)
                            st.rerun()
                    st.markdown("---")
                    
        st.markdown("### Tugas Eksekusi Ditolak (Perlu Revisi)")
        if rej.empty: 
            st.info("Catatan evaluasi sempurna. Tidak ada instruksi pengerjaan yang memerlukan revisi lanjutan.")
        else:
            for _, row in rej.iterrows():
                with st.container():
                    st.markdown(f"**Tugas Membutuhkan Koreksi:** {row['task_name']}")
                    
                    c1, c2 = st.columns([1, 4])
                    with c1:
                        pg = pg_latest(row["task_id"])
                        if pg:
                            foto_show(pg.get("photo_report",""), "Bukti Pengajuan Lama", use_col_width=True)
                    with c2:
                        st.error(f"Catatan Arahan Revisi: {row.get('tl_comment', 'Tidak terdapat alasan rinci.')}")

                    cam = st.camera_input("Lampirkan Bukti Revisi Aktual", key=f"rj_{row['task_id']}")
                    
                    if st.button(f"Kirim Revisi Laporan: {row['task_name']}", key=f"rb_{row['task_id']}", use_container_width=True):
                        if not cam: 
                            st.error("Protokol Audit: Resolusi visual hasil revisi wajib dilampirkan kembali.")
                        else:
                            with st.spinner("Memproses transmisi pembaruan dokumen..."): 
                                url = foto_compress(cam)
                                pg_save(row['task_id'], eid, st.session_state.user_name, url)
                            st.success("Bukti perbaikan telah berhasil diajukan kembali ke tinjauan Pengawas.")
                            time.sleep(1.5)
                            st.rerun()
                    st.markdown("---")
                    
        st.markdown("### Menunggu Verifikasi Pemeriksaan Pengawas")
        if pend.empty: 
            st.info("Pusat antrean bersih. Tidak terdapat laporan aktif yang sedang menunggu proses peninjauan.")
        else:
            for _, row in pend.iterrows():
                with st.container():
                    st.warning(f"Tugas: {row['task_name']} - Laporan berada dalam antrean proses evaluasi Pengawas Divisi.")
                    pg = pg_latest(row['task_id'])
                    if pg:
                        c1, c2 = st.columns([1, 4])
                        with c1:
                            foto_show(pg.get("photo_report", ""), "Lampiran Diajukan", use_col_width=True)
                        with c2:
                            st.write("**Status Dokumen**")
                            st.info("Sedang ditinjau oleh Administrator / Pengawas Lapangan.")
                    st.markdown("---")
                    
        st.markdown("### Riwayat Selesai & Disetujui")
        if appr.empty: 
            st.info("Arsip historis kosong. Belum terdapat laporan tugas yang diverifikasi sebagai Selesai.")
        else:
            for _, row in appr.iterrows():
                with st.container():
                    st.success(f"Disetujui: **{row['task_name']}**")
                    pg = pg_latest(row['task_id'])
                    if pg:
                        c1, c2 = st.columns([1, 4])
                        with c1:
                            foto_show(pg.get("photo_report",""), "Lampiran Final", use_col_width=True)
                        with c2:
                            st.write("**Resolusi Pengawas**")
                            if row.get("tl_comment"):
                                st.caption(f"Catatan Evaluasi: {row['tl_comment']}")
                            else:
                                st.caption("Tugas divalidasi tanpa catatan tambahan.")
                    st.markdown("---")
                    
        st.markdown("---")
        st.subheader("Papan Status Kinerja Personel Divisi")
        da = tk_all(eid=eid, div=mydiv) 
        if da.empty: 
            st.info("Sistem pemantauan belum merekam riwayat pergerakan alokasi kerja untuk divisi operasional ini.")
        else: 
            st.dataframe(da[['crew_name', 'task_name', 'status']], use_container_width=True)